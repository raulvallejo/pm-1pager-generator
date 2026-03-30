import os
import io
import re
import html
import time
import json
import uuid
from fastapi import FastAPI, APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from langchain_anthropic import ChatAnthropic
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from tavily import TavilyClient
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import opik
from opik.api_objects.prompt.client import PromptClient as _PromptClient
from opik.api_objects import opik_client as _opik_client_module

# ---------------------------------------------------------------------------
# Load environment variables from backend/.env
# ANTHROPIC_API_KEY, TAVILY_API_KEY, OPIK_API_KEY etc. are read from here
# in local dev. On Render they are set directly in the dashboard.
# NOTE: opik.configure() is never called — OPIK reads its config from env
# vars (OPIK_API_KEY, OPIK_PROJECT_NAME, OPIK_WORKSPACE, OPIK_URL_OVERRIDE)
# automatically at import time, which avoids the interactive TTY prompt.
# ---------------------------------------------------------------------------
load_dotenv()

print("OPIK initialized via env vars")


# ---------------------------------------------------------------------------
# Safe OPIK tracking decorator
#
# Wraps opik.track() so that any failure at decoration time (bad API key,
# network issue, version incompatibility) logs a warning and falls back to
# the plain function rather than crashing the server.
#
# Usage:  @_safe_track(name="my_span")
# ---------------------------------------------------------------------------
def _safe_track(name: str, type: str = "general"):
    def decorator(func):
        try:
            return opik.track(name=name, type=type)(func)
        except Exception as e:
            print(f"WARNING: OPIK @track setup failed for '{name}': {e} — running untracked")
            return func
    return decorator

app = FastAPI(title="PM 1-Pager Generator API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------------------------------------------------------
# API router
#
# All functional endpoints are mounted under /api so the frontend's
# production calls (e.g. https://…onrender.com/api/chat) match exactly.
# The health check stays on `app` directly at GET / so Render's health
# probe works without any prefix.
# ---------------------------------------------------------------------------
router = APIRouter(prefix="/api")

# ---------------------------------------------------------------------------
# Claude client via LangChain
#
# ChatAnthropic reads ANTHROPIC_API_KEY from the environment automatically.
# One shared instance at startup — thread-safe.
# ---------------------------------------------------------------------------
llm = ChatAnthropic(model="claude-haiku-4-5")

# ---------------------------------------------------------------------------
# Tavily client
#
# TavilyClient wraps the Tavily Search REST API.
# .search(query, max_results) returns a dict with a "results" list, where
# each result has "title", "url", and "content" fields.
# ---------------------------------------------------------------------------
tavily_client = TavilyClient(api_key=os.getenv("TAVILY_API_KEY"))

# ---------------------------------------------------------------------------
# Exponential backoff helper
#
# Retries the Claude call up to MAX_RETRIES times on rate limit errors.
# Waits 2^attempt seconds between tries: 2s, 4s, 8s.
# ---------------------------------------------------------------------------
MAX_RETRIES = 3

def invoke_with_backoff(messages: list) -> str:
    for attempt in range(MAX_RETRIES):
        try:
            return llm.invoke(messages).content
        except Exception as e:
            err = str(e).lower()
            is_rate_limit = "quota" in err or "429" in err or "rate" in err
            if is_rate_limit and attempt < MAX_RETRIES - 1:
                wait = 2 ** (attempt + 1)  # 2s, 4s, 8s
                time.sleep(wait)
                continue
            raise


def _invoke_with_usage(messages: list) -> tuple[str, dict]:
    """invoke_with_backoff + token usage and latency capture."""
    for attempt in range(MAX_RETRIES):
        try:
            t0 = time.time()
            result = llm.invoke(messages)
            latency_ms = round((time.time() - t0) * 1000)

            usage = getattr(result, "usage_metadata", {}) or {}
            input_tokens  = usage.get("input_tokens", 0)
            output_tokens = usage.get("output_tokens", 0)
            total_tokens  = usage.get("total_tokens", input_tokens + output_tokens)
            estimated_cost = (input_tokens / 1_000_000 * 1.00) + (output_tokens / 1_000_000 * 5.00)

            return result.content, {
                "prompt_tokens":      input_tokens,    # OpenAI key names required by OPIK UI
                "completion_tokens":  output_tokens,
                "total_tokens":       total_tokens,
                "latency_ms":         latency_ms,
                "estimated_cost_usd": round(estimated_cost, 8),
            }
        except Exception as e:
            err = str(e).lower()
            is_rate_limit = "quota" in err or "429" in err or "rate" in err
            if is_rate_limit and attempt < MAX_RETRIES - 1:
                wait = 2 ** (attempt + 1)
                time.sleep(wait)
                continue
            raise

# ---------------------------------------------------------------------------
# Session stores
#
# sessions       — conversation history per session_id (list of role/content dicts)
# session_research — research summary per session_id (populated by /research)
#
# Both are plain dicts keyed by the UUID the frontend generates.
# They disappear on server restart (Sprint 4 can add a real DB).
# ---------------------------------------------------------------------------
sessions: dict[str, list[dict]] = {}
session_research: dict[str, str] = {}

# session_documents stores the raw 1-pager markdown text per session.
# Populated whenever a completed 1-pager is returned (from /chat or /research).
# The /download/* endpoints read from here to generate files.
session_documents: dict[str, str] = {}

# ---------------------------------------------------------------------------
# PII Anonymizer
#
# Strips personally identifiable information from user text.
# The LLM still receives the original unsanitized text so response quality is
# completely unaffected.
#
# Patterns detected:
#   Email addresses — john@example.com          → [EMAIL]
#   Phone numbers   — +1 (555) 123-4567         → [PHONE]
#   Names           — context-triggered only    → [NAME]
#                     (e.g. "My name is John Smith", "I'm Sarah")
#
# WHY context-triggered for names?
#   A naive "two consecutive Title-Case words" regex would fire on product
#   content like "Problem Statement", "Market Context", "Key Metrics" —
#   every section header in a 1-pager. By requiring a lead phrase like
#   "my name is" or "I'm", we keep precision high and false positives low.
# ---------------------------------------------------------------------------

# Matches standard email addresses
_EMAIL_RE = re.compile(
    r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b'
)

# Matches US/international phone numbers in common formats:
#   +1 (555) 123-4567 | 555-123-4567 | 555.123.4567 | 5551234567
_PHONE_RE = re.compile(
    r'\b(?:\+?1[\s.\-]?)?'        # optional +1 country code
    r'(?:\(?\d{3}\)?[\s.\-]?)'    # area code, optional parens
    r'\d{3}[\s.\-]?\d{4}\b'       # 7-digit subscriber number
)

# Only fires when a name is introduced with a known trigger phrase.
# Captures 1–3 Title-Case word sequences following the phrase.
_NAME_RE = re.compile(
    r"(?i)"                                                       # case-insensitive
    r"(?:my name is|i'?m|i am|call me|this is|contact|by)\s+"   # trigger phrases
    r"([A-Z][a-z]{1,20}(?:\s+[A-Z][a-z]{1,20}){0,2})"          # 1–3 name parts
)


def anonymize_pii(text: str) -> str:
    """Replace PII patterns in text with safe placeholders."""
    if not isinstance(text, str) or not text:
        return text
    text = _EMAIL_RE.sub("[EMAIL]", text)
    text = _PHONE_RE.sub("[PHONE]", text)
    # Substitute only the captured name group, preserving the trigger phrase
    text = _NAME_RE.sub(lambda m: m.group(0).replace(m.group(1), "[NAME]"), text)
    return text


def anonymize_history(history: list) -> list:
    """Return a sanitized copy of a conversation history list."""
    return [
        {"role": msg["role"], "content": anonymize_pii(msg["content"])}
        for msg in history
    ]


# ---------------------------------------------------------------------------
# System prompt
#
# Changed from Sprint 2: step 4 now tells Claude to output [READY_FOR_RESEARCH]
# instead of generating the 1-pager directly. This lets the backend intercept,
# run Tavily searches, inject the results, and THEN ask Claude for the 1-pager
# with real market data baked in.
#
# The 1-pager format is kept here so Claude knows what it will eventually produce.
# A new "Market Context" section uses the injected Tavily research.
# ---------------------------------------------------------------------------
SYSTEM_PROMPT = """\
You are a Product Management assistant that helps teams turn rough ideas into
structured PM 1-pagers.

WORKFLOW
--------
1. The user describes a product initiative (may be vague or detailed).
2. You analyse it and identify the single most important piece of missing
   information.
3. You ask ONE clarifying question per response — maximum 3 questions total.
4. After at most 3 questions (or sooner if the description is already
   detailed), respond ONLY with the exact token [READY_FOR_RESEARCH] on its
   own line — nothing else. The system will automatically research the market
   and then ask you to generate the 1-pager with that context.

RULES
-----
- Ask EXACTLY ONE question per turn. Never bundle multiple questions.
- After 3 questions you MUST output [READY_FOR_RESEARCH], even if you wish
  you had more information.
- If the user's initial message is already detailed enough, skip questions
  and output [READY_FOR_RESEARCH] immediately.
- NEVER combine a question and [READY_FOR_RESEARCH] in the same response.

1-PAGER FORMAT
--------------
When you are asked to produce the 1-pager (after research context is injected),
your ENTIRE response must be the following markdown:

---
## PM 1-Pager: [Initiative Name]

### Problem Statement
[What problem are we solving, for whom, and why does it matter now?]

### Target User
[Specific user persona — role, context, key pain point]

### Proposed Solution
[What we are building and the core user experience in 2-3 sentences]

### Key Metrics
[2–4 measurable success indicators; include targets where possible]

### Market Context
[2–3 sentences drawing on the provided research: market size, key trends,
 and notable competitors]

### Risks & Assumptions
[Top 3 risks or assumptions that could invalidate this initiative]
---

IMPORTANT: start your response with the exact line "---" (three dashes, nothing
else on that line) so the frontend can detect and render it as a document.
"""


def _fetch_opik_prompt(name: str) -> str:
    """Fetch latest prompt template from OPIK Prompt Library. Falls back to SYSTEM_PROMPT on any failure."""
    try:
        client = _opik_client_module.get_client_cached()
        prompt_client = _PromptClient(client.rest_client)
        version = prompt_client.get_prompt(name=name)
        if version and version.template:
            print(f"Loaded '{name}' from OPIK Prompt Library (commit: {version.commit})")
            return version.template
        print(f"WARNING: OPIK prompt '{name}' not found — using hardcoded fallback")
    except Exception as e:
        print(f"WARNING: Could not fetch '{name}' from OPIK Prompt Library: {e} — using hardcoded fallback")
    return SYSTEM_PROMPT


ACTIVE_SYSTEM_PROMPT = _fetch_opik_prompt("1pager-system-prompt")


# ---------------------------------------------------------------------------
# Helper: convert stored history dicts → LangChain message objects
#
# The Anthropic API is stateless — every call must include the full history.
# LangChain needs typed objects (HumanMessage / AIMessage), not raw dicts.
# We factor this out so both /chat and /research can reuse it.
# ---------------------------------------------------------------------------
def build_lc_messages(history: list) -> list:
    result = []
    for msg in history:
        if msg["role"] == "user":
            result.append(HumanMessage(content=msg["content"]))
        else:
            result.append(AIMessage(content=msg["content"]))
    return result


# ---------------------------------------------------------------------------
# Helper: parse the 1-pager markdown into a structured dict
#
# The 1-pager always follows the format:
#   ---
#   ## PM 1-Pager: Title
#   ### Section Name
#   content...
#   ---
#
# Returns: { "title": str, "sections": [{"heading": str, "content": str}] }
#
# Both generate_docx and generate_pdf call this so the parsing logic lives
# in one place — change the format here and both output types update.
# ---------------------------------------------------------------------------
def parse_1pager(text: str) -> dict:
    lines = text.strip().split("\n")
    # Strip leading and trailing "---" delimiter lines
    if lines and lines[0].strip() == "---":
        lines = lines[1:]
    if lines and lines[-1].strip() == "---":
        lines = lines[:-1]

    result = {"title": "PM 1-Pager", "sections": []}
    current_heading = None
    current_content: list[str] = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("## "):
            result["title"] = stripped[3:].strip()
        elif stripped.startswith("### "):
            # Save the previous section before starting the new one
            if current_heading is not None:
                result["sections"].append({
                    "heading": current_heading,
                    "content": "\n".join(current_content).strip(),
                })
            current_heading = stripped[4:].strip()
            current_content = []
        else:
            if current_heading is not None:
                current_content.append(line)

    # Don't forget the last section
    if current_heading is not None:
        result["sections"].append({
            "heading": current_heading,
            "content": "\n".join(current_content).strip(),
        })

    return result


# ---------------------------------------------------------------------------
# Helper: generate a .docx file from parsed 1-pager data
#
# python-docx builds Word documents programmatically:
#   - add_heading(text, level=1) → large bold heading (maps to Word's H1/H2)
#   - add_paragraph(text)        → normal body text
#   - doc.save(buffer)           → writes the .docx bytes into the buffer
#
# We use an io.BytesIO buffer so the file never touches disk — the bytes
# go straight into the HTTP response.
# ---------------------------------------------------------------------------
def generate_docx(parsed: dict) -> bytes:
    doc = Document()

    # Title — Heading 1
    title_para = doc.add_heading(parsed["title"], level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for section in parsed["sections"]:
        # Section header — Heading 2
        doc.add_heading(section["heading"], level=2)
        # Section body — Normal paragraph
        # Add a paragraph per non-empty line so spacing looks clean
        for line in section["content"].split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Helper: generate a .pdf file from parsed 1-pager data
#
# reportlab builds PDFs from a "story" — a list of flowable objects
# (Paragraph, Spacer, etc.) that are laid out top-to-bottom on the page.
#
# Key concepts:
#   ParagraphStyle — defines font, size, spacing for a class of text
#   Paragraph(text, style) — a block of styled text; accepts basic HTML tags
#   Spacer(width, height)  — blank vertical space between elements
#   SimpleDocTemplate.build(story) — renders the story to the buffer
#
# We call html.escape() on all user-supplied content because reportlab
# parses Paragraph text as XML — unescaped < > & would crash the build.
# ---------------------------------------------------------------------------
def generate_pdf(parsed: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    base_styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "DocTitle",
        parent=base_styles["Title"],
        fontSize=20,
        textColor=colors.HexColor("#1a1a1a"),
        spaceAfter=16,
    )
    heading_style = ParagraphStyle(
        "DocHeading",
        parent=base_styles["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#2563eb"),
        spaceBefore=14,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "DocBody",
        parent=base_styles["Normal"],
        fontSize=11,
        leading=17,
        textColor=colors.HexColor("#1a1a1a"),
        spaceAfter=4,
    )

    story = []
    story.append(Paragraph(html.escape(parsed["title"]), title_style))
    story.append(Spacer(1, 0.3 * cm))

    for section in parsed["sections"]:
        story.append(Paragraph(html.escape(section["heading"]), heading_style))
        for line in section["content"].split("\n"):
            if line.strip():
                story.append(Paragraph(html.escape(line.strip()), body_style))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Helper: rule-based eval scoring for the generated 1-pager
#
# Returns a dict of three scores, each 0.0–1.0:
#
#   completeness   — fraction of the 6 required sections present.
#                    1.0 = all sections exist, 0.0 = empty document.
#
#   research_usage — 1.0 if the Market Context section has substantial
#                    content (>50 chars), meaning Claude actually used the
#                    Tavily research data. 0.0 if the section is missing
#                    or nearly empty.
#
#   clarity        — grades the Problem Statement on length as a proxy for
#                    specificity: 1.0 ≥100 chars, 0.5 ≥30 chars, 0.0 otherwise.
#
# These scores are logged to OPIK as feedback on the 1pager_generation span,
# giving you a per-trace quality signal visible in the OPIK dashboard.
# ---------------------------------------------------------------------------
def score_1pager(text: str) -> dict[str, float]:
    parsed = parse_1pager(text)
    sections = {s["heading"]: s["content"] for s in parsed["sections"]}

    required_sections = {
        "Problem Statement", "Target User", "Proposed Solution",
        "Key Metrics", "Market Context", "Risks & Assumptions",
    }
    present = sum(1 for h in required_sections if h in sections)
    completeness = round(present / len(required_sections), 2)

    market_content = sections.get("Market Context", "")
    research_usage = 1.0 if len(market_content) > 50 else 0.0

    problem_content = sections.get("Problem Statement", "")
    if len(problem_content) >= 100:
        clarity = 1.0
    elif len(problem_content) >= 30:
        clarity = 0.5
    else:
        clarity = 0.0

    return {
        "completeness": completeness,
        "research_usage": research_usage,
        "clarity": clarity,
    }


# ---------------------------------------------------------------------------
# Helper: run Tavily market research for a session
#
# Flow:
#   1. Ask Claude to name the initiative in 3-5 words (used as query base).
#   2. Build 3 targeted search queries: market size, trends, competitors.
#   3. Call Tavily for each query and collect the top 3 snippets per query.
#   4. Return a formatted research string to inject into the 1-pager prompt.
#
# Why ask Claude for the initiative name?
#   The conversation may span multiple turns. Extracting the name via Claude
#   is more reliable than using the raw first message.
#
# Why 3 queries?
#   Each covers a different angle (size, trends, competition) that maps
#   directly to the new "Market Context" section of the 1-pager.
# ---------------------------------------------------------------------------
def research_initiative(session_id: str) -> str:
    history = sessions[session_id]

    # Step 1: ask Claude for a concise initiative name to build queries from
    name_messages = [
        SystemMessage(content=(
            "Read this product conversation and return a concise 3-5 word name "
            "for the initiative being discussed. Return ONLY the name — no "
            "punctuation, no explanation, nothing else."
        )),
        *build_lc_messages(history),
    ]
    initiative_name = invoke_with_backoff(name_messages).strip().strip(".")

    # Step 2: build 3 targeted search queries
    queries = [
        f"{initiative_name} market size",
        f"{initiative_name} industry trends 2025",
        f"{initiative_name} top competitors",
    ]

    # Step 3: run Tavily searches and collect snippets
    lines = [f"MARKET RESEARCH FOR: {initiative_name}\n"]
    for query in queries:
        lines.append(f"\n### {query}")
        try:
            results = tavily_client.search(query=query, max_results=3)
            for r in results.get("results", [])[:3]:
                # Truncate long content to keep the injected context concise
                snippet = r.get("content", "")[:250].replace("\n", " ")
                lines.append(f"- {r['title']}: {snippet}")
        except Exception:
            lines.append("- (search unavailable for this query)")

    return "\n".join(lines)


@_safe_track(name="clarification_questions")
def track_clarification(session_id: str, message: str, history: list) -> str:
    lc_messages = [SystemMessage(content=ACTIVE_SYSTEM_PROMPT)] + build_lc_messages(history)
    return invoke_with_backoff(lc_messages)


def track_web_research(session_id: str) -> str:
    return research_initiative(session_id)


@_safe_track(name="1pager_generation", type="llm")
def track_1pager_generation(session_id: str, history: list, research_summary: str) -> str:
    lc_messages = [SystemMessage(content=ACTIVE_SYSTEM_PROMPT)]
    lc_messages.extend(build_lc_messages(history))
    lc_messages.append(HumanMessage(content=(
        f"Research complete. Here is the market context to use when writing "
        f"the Market Context section of the 1-pager:\n\n"
        f"{research_summary}\n\n"
        f"Now generate the PM 1-pager."
    )))
    content, usage = _invoke_with_usage(lc_messages)
    try:
        opik.opik_context.update_current_span(
            usage={
                "prompt_tokens":     usage["prompt_tokens"],
                "completion_tokens": usage["completion_tokens"],
                "total_tokens":      usage["total_tokens"],
                "total_cost":        usage["estimated_cost_usd"],
            },
            metadata={"latency_ms": usage["latency_ms"]},
        )
    except Exception as e:
        print(f"WARNING: OPIK span metadata update failed: {e}")
    return content


@_safe_track(name="1pager_pipeline")
def generate_1pager_pipeline(session_id: str) -> str:
    """
    TOP-LEVEL TRACE: the full research → generation pipeline.

    This is the function the OPIK dashboard shows as the root trace entry.
    It calls the two child spans in sequence:
      1. track_web_research  — 3 Tavily searches
      2. track_1pager_generation — Claude generates 1-pager + eval scores logged

    Called by the /research endpoint. All side effects (updating sessions,
    session_research, session_documents) remain in the endpoint; this function
    is pure pipeline logic.
    """
    history = sessions[session_id]

    # Child span 1 — web research (graceful degradation if Tavily fails)
    try:
        research_summary = track_web_research(session_id)
    except Exception:
        research_summary = (
            "(Web research unavailable — generating from conversation context only.)"
        )

    # Store research summary so the session retains it
    session_research[session_id] = research_summary

    # Child span 2 — 1-pager generation + eval scoring
    reply = track_1pager_generation(session_id, history, research_summary)

    return reply, research_summary


# ---------------------------------------------------------------------------
# Request / Response models
# ---------------------------------------------------------------------------
class ChatRequest(BaseModel):
    session_id: str   # UUID generated by the frontend
    message: str      # The user's latest message


class ChatResponse(BaseModel):
    reply: str         # Claude's response text
    session_id: str    # Echoed back
    is_complete: bool  # True when the 1-pager has been generated
    is_researching: bool  # True when Claude signalled [READY_FOR_RESEARCH]


class ResearchRequest(BaseModel):
    session_id: str   # Must match an existing session


class ResearchResponse(BaseModel):
    reply: str         # The generated 1-pager
    session_id: str
    is_complete: bool  # Always True when research succeeds


# ---------------------------------------------------------------------------
# /chat endpoint (unchanged clarification flow + [READY_FOR_RESEARCH] detection)
#
# New in Sprint 3:
#   - When Claude replies with [READY_FOR_RESEARCH], we do NOT store that
#     token in history (it's an internal signal, not a real assistant turn).
#     Instead we return is_researching=True and a user-friendly message.
#   - The frontend sees is_researching=True, shows "Researching..." in the
#     chat, and immediately fires a POST /research call.
# ---------------------------------------------------------------------------
@router.post("/chat", response_model=ChatResponse)
def chat(request: ChatRequest):
    # 1. Get or create session history
    if request.session_id not in sessions:
        sessions[request.session_id] = []

    history = sessions[request.session_id]

    # 2. Append user message
    history.append({"role": "user", "content": request.message})

    # 3. Call Claude via the OPIK-tracked wrapper.
    #    track_clarification() creates a "clarification_questions" trace in
    #    OPIK and records the input message + output reply automatically.
    try:
        reply = track_clarification(request.session_id, request.message, history)
    except Exception as e:
        err = str(e).lower()
        if "api key" in err or "credential" in err or "401" in err or "authentication" in err:
            raise HTTPException(status_code=401, detail="Invalid ANTHROPIC_API_KEY.")
        if "quota" in err or "429" in err or "rate" in err or "overloaded" in err:
            raise HTTPException(status_code=429, detail="Anthropic rate limit hit — all retries exhausted.")
        raise HTTPException(status_code=503, detail=f"Could not reach Anthropic API: {e}")

    # 4. Check if Claude is signalling it's ready for research
    #    We store a clean placeholder in history instead of the raw token so
    #    the conversation context stays coherent for the /research call.
    if "[READY_FOR_RESEARCH]" in reply:
        history.append({"role": "assistant", "content": "[READY_FOR_RESEARCH]"})
        return ChatResponse(
            reply="Researching the market for you...",
            session_id=request.session_id,
            is_complete=False,
            is_researching=True,
        )

    # 5. Normal question turn — store reply and return
    history.append({"role": "assistant", "content": reply})
    is_complete = reply.strip().startswith("---")
    # Cache 1-pager text so /download/* endpoints can retrieve it later
    if is_complete:
        session_documents[request.session_id] = reply
    return ChatResponse(
        reply=reply,
        session_id=request.session_id,
        is_complete=is_complete,
        is_researching=False,
    )


# ---------------------------------------------------------------------------
# /research endpoint (Sprint 3 — runs Tavily + generates enriched 1-pager)
#
# Called automatically by the frontend immediately after receiving
# is_researching=True from /chat. Flow:
#   1. Run research_initiative() — asks Claude for a name, runs 3 Tavily
#      searches, returns a formatted research summary string.
#   2. Inject the summary into the LangChain message list as a HumanMessage
#      so Claude sees it as "provided context" in the conversation.
#   3. Call Claude to generate the 1-pager using that context.
#   4. Append the 1-pager to session history and return it.
#
# Why inject as HumanMessage and not SystemMessage?
#   Anthropic's API only allows one SystemMessage. We've already used it for
#   SYSTEM_PROMPT, so we pass research data as a HumanMessage at the end of
#   the conversation — Claude treats it as the final piece of input before
#   generating the document.
# ---------------------------------------------------------------------------
@router.post("/research", response_model=ResearchResponse)
def research(request: ResearchRequest):
    if request.session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found. Start a conversation first.")

    history = sessions[request.session_id]

    # Run the full OPIK-traced pipeline:
    #   generate_1pager_pipeline() is the top-level trace in OPIK.
    #   Inside it, track_web_research() and track_1pager_generation() are
    #   child spans. Eval scores (completeness, research_usage, clarity) are
    #   logged automatically onto the 1pager_generation span.
    try:
        reply, research_summary = generate_1pager_pipeline(request.session_id)
    except Exception as e:
        err = str(e).lower()
        if "api key" in err or "credential" in err or "authentication" in err:
            raise HTTPException(status_code=401, detail="Invalid ANTHROPIC_API_KEY.")
        if "quota" in err or "429" in err or "rate" in err or "overloaded" in err:
            raise HTTPException(status_code=429, detail="Anthropic rate limit hit.")
        raise HTTPException(status_code=503, detail=f"Could not reach Anthropic API: {e}")

    # Append the final 1-pager to history and cache it for downloads
    history.append({"role": "assistant", "content": reply})
    session_documents[request.session_id] = reply

    return ResearchResponse(
        reply=reply,
        session_id=request.session_id,
        is_complete=reply.strip().startswith("---"),
    )


# ---------------------------------------------------------------------------
# Download endpoints (Sprint 4)
#
# Both endpoints follow the same pattern:
#   1. Look up the cached 1-pager text for the session.
#   2. Parse it into sections with parse_1pager().
#   3. Generate the file bytes in memory (no disk I/O).
#   4. Return a StreamingResponse with the correct MIME type and a
#      Content-Disposition header so the browser triggers a file download.
#
# Content-Disposition: attachment; filename="..." tells the browser to save
# the response as a file rather than trying to display it inline.
# ---------------------------------------------------------------------------
class DownloadRequest(BaseModel):
    session_id: str


@router.post("/download/docx")
def download_docx(request: DownloadRequest):
    if request.session_id not in session_documents:
        raise HTTPException(
            status_code=404,
            detail="No 1-pager found for this session. Generate one first.",
        )
    parsed = parse_1pager(session_documents[request.session_id])
    docx_bytes = generate_docx(parsed)
    # Derive a clean filename from the initiative title
    safe_name = parsed["title"].replace(" ", "_").replace("/", "-")[:60]
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
    )


@router.post("/download/pdf")
def download_pdf(request: DownloadRequest):
    if request.session_id not in session_documents:
        raise HTTPException(
            status_code=404,
            detail="No 1-pager found for this session. Generate one first.",
        )
    parsed = parse_1pager(session_documents[request.session_id])
    pdf_bytes = generate_pdf(parsed)
    safe_name = parsed["title"].replace(" ", "_").replace("/", "-")[:60]
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.pdf"'},
    )


# ---------------------------------------------------------------------------
# Legacy /generate endpoint (Sprint 1 — kept so nothing breaks)
# ---------------------------------------------------------------------------
class GenerateRequest(BaseModel):
    message: str


class GenerateResponse(BaseModel):
    reply: str


@router.post("/generate", response_model=GenerateResponse)
def generate(request: GenerateRequest):
    """Sprint 1 mock endpoint — still works but /chat is the real one now."""
    mock_reply = (
        f"[Legacy /generate] Got: \"{request.message}\". "
        "Please use the /chat endpoint instead."
    )
    return GenerateResponse(reply=mock_reply)


# ---------------------------------------------------------------------------
# Health check
# ---------------------------------------------------------------------------
@app.api_route("/", methods=["GET", "HEAD"])
def health_check():
    return {"status": "ok"}


@app.get("/api/test")
def test():
    return {"message": "api prefix working", "version": "sprint5"}


@app.get("/api/chat/test")
def chat_test():
    return {"message": "chat route exists", "version": "sprint5"}


# Must come AFTER all @router.post/@router.get definitions above.
# include_router snapshots the router's route list at call time —
# any routes defined after this line would be silently ignored.
app.include_router(router)
